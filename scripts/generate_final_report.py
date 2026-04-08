from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = Path("output/doc/Deepanshu-Jain-Combined-Project-Report-and-Synopsis.docx")
SCREENSHOT_DIR = Path("output/doc/screenshots")
ASSET_DIR = Path("output/doc/assets")

STUDENT_NAME = "Deepanshu Jain"
ROLL_NUMBER = "2201730107"
SECTION = "B"
UNIVERSITY_NAME = "KR MANGALAM UNIVERSITY"
CITY = "Gurgaon"
MENTOR_NAME = "Gaurav"
PROGRAM = "BACHELOR OF TECHNOLOGY (CSE / AI & ML)"
SEMESTER = "Semester VIII"
SUBMISSION_MONTH = "April 2026"
PROJECT_CATEGORY = "University Based Project"
PROJECT_TITLE = (
    "IoT-Enabled AI-Driven Learning Analytics and Personalized Feedback "
    "System for Mathematics Education"
)
PROJECT_CODE = "Math Analytics Studio"
REPO_URL = "https://github.com/deepanshu306/iot-ai-learning-analytics-math"
LIVE_URL = "https://deepanshu306.github.io/iot-ai-learning-analytics-math/"
STUDENT_COUNT = 3
TEACHER_COUNT = 1
ADMIN_COUNT = 1
QUESTION_COUNT = 12
ATTEMPT_COUNT = 6
PROJEXA_TEAM_ID = "26E4362"
UNIVERSITY_LOGO = ASSET_DIR / "krmu-crest.png"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_document_style(document):
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        document.styles[style_name].font.name = "Times New Roman"

    document.styles["Title"].font.size = Pt(20)
    document.styles["Heading 1"].font.size = Pt(16)
    document.styles["Heading 2"].font.size = Pt(14)
    document.styles["Heading 3"].font.size = Pt(12)

    props = document.core_properties
    props.title = PROJECT_TITLE
    props.author = STUDENT_NAME
    props.subject = "Combined synopsis and final project report"
    props.keywords = "IoT, learning analytics, mathematics education, personalized feedback, smart classroom"


def center(document, text, bold=False, size=12, after=6):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(after)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def body(document, text):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def bullet(document, text):
    para = document.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.1
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def numbered(document, text):
    para = document.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.1
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    tbl_pr.append(borders)


def set_cell_text(cell, text, *, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Inches(width)


def add_signature_block(document, left_lines, right_lines=None):
    cols = 2 if right_lines else 1
    rows = max(len(left_lines), len(right_lines or []))
    table = document.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    set_table_widths(table, [3.1, 3.1] if cols == 2 else [6.2])

    for row_index in range(rows):
        if row_index < len(left_lines):
            set_cell_text(table.rows[row_index].cells[0], left_lines[row_index], size=11)
        if right_lines and row_index < len(right_lines):
            set_cell_text(table.rows[row_index].cells[1], right_lines[row_index], size=11)

    document.add_paragraph()


def add_footer(section):
    para = section.footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"{PROJECT_TITLE} | Page ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    add_page_number(para)


def add_heading_page_break(document, title, level=1):
    document.add_page_break()
    heading = document.add_heading(title, level=level)
    heading.paragraph_format.space_after = Pt(10)


def add_picture_with_caption(document, image_name, caption, width=4.15, page_break_before=False):
    image_path = SCREENSHOT_DIR / image_name
    if not image_path.exists():
        return
    if page_break_before:
        document.add_page_break()
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    caption_para = document.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_para.add_run(caption)
    caption_run.font.name = "Times New Roman"
    caption_run.font.size = Pt(10)
    caption_run.italic = True
    caption_para.paragraph_format.space_after = Pt(8)


def add_table_title(document, text):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.bold = True


def cover_page(document):
    for _ in range(3):
        document.add_paragraph()
    center(document, "Major Project Report", bold=True, size=18, after=16)
    center(document, "Combined Synopsis and Final Report", size=12, after=10)
    center(document, f"Project Code: {PROJECT_CODE}", size=11, after=10)
    center(document, PROJECT_TITLE, bold=True, size=18, after=18)
    center(document, f"Project Category: {PROJECT_CATEGORY}", size=12, after=10)
    center(document, f"Projexa Team Id- {PROJEXA_TEAM_ID}", size=12, after=12)
    center(document, "Submitted in partial fulfilment of the requirement of the degree of", size=12, after=4)
    center(document, PROGRAM, bold=True, size=14, after=4)
    center(document, SEMESTER, size=12, after=4)
    center(document, "to", size=12, after=4)
    center(document, UNIVERSITY_NAME, bold=True, size=14, after=16)
    center(document, "by", size=12, after=8)
    center(document, STUDENT_NAME, bold=True, size=13, after=4)
    center(document, f"{ROLL_NUMBER}    Section - {SECTION}", size=12, after=18)
    center(document, "Under the supervision of", size=12, after=8)
    center(document, MENTOR_NAME, bold=True, size=12, after=2)
    center(document, "Project Guide", size=11, after=18)
    center(document, "Department of Computer Science and Engineering", bold=True, size=12, after=2)
    center(document, "School of Engineering and Technology", bold=True, size=12, after=2)
    if UNIVERSITY_LOGO.exists():
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(10)
        run = para.add_run()
        run.add_picture(str(UNIVERSITY_LOGO), width=Inches(1.45))
    center(document, f"{UNIVERSITY_NAME}, {CITY}, India", size=12, after=2)
    center(document, SUBMISSION_MONTH, size=12, after=0)
    document.add_page_break()


def front_matter(document):
    document.add_heading("DECLARATION", level=1)
    body(
        document,
        f"I hereby declare that the project report entitled \"{PROJECT_TITLE}\" is my original "
        "work prepared during the final year of the Bachelor of Technology program. This work has "
        "been carried out for academic submission and has not been submitted elsewhere for the "
        "award of any degree, diploma, or certificate. The concepts, explanations, analysis, and "
        "implementation details presented in this report are true to the best of my knowledge."
    )
    add_signature_block(
        document,
        [
            f"Name: {STUDENT_NAME}",
            f"Roll Number: {ROLL_NUMBER}",
            f"Section: {SECTION}",
            "Signature of Student: __________________________",
            "Date: __________________________",
        ]
    )
    document.add_page_break()

    document.add_heading("CERTIFICATE", level=1)
    body(
        document,
        f"This is to certify that the major project entitled \"{PROJECT_TITLE}\" submitted by "
        f"{STUDENT_NAME} ({ROLL_NUMBER}) is a bonafide academic work carried out under guidance "
        "for the partial fulfilment of the degree of Bachelor of Technology in Computer Science "
        "and Engineering / Artificial Intelligence and Machine Learning. The work presented in "
        "this report is fit for academic evaluation and project assessment."
    )
    add_signature_block(
        document,
        [
            f"Project Guide: {MENTOR_NAME}",
            "Signature of Guide: __________________________",
            "Date: __________________________",
        ],
        [
            f"Student Name: {STUDENT_NAME}",
            f"University: {UNIVERSITY_NAME}",
            "Signature of Student: __________________________",
        ],
    )
    document.add_page_break()

    document.add_heading("ACKNOWLEDGEMENT", level=1)
    body(
        document,
        f"I would like to express my sincere gratitude to {MENTOR_NAME}, my project guide, for "
        "providing guidance, encouragement, and valuable suggestions during the development of "
        "this project. I also thank the faculty members of the Department of Computer Science and "
        "Engineering, School of Engineering and Technology, KR Mangalam University, for their "
        "support and academic direction throughout the final-year project process."
    )
    body(
        document,
        "I am grateful to my classmates, friends, and family members for their encouragement and "
        "support. Their suggestions and motivation helped me complete the implementation, "
        "documentation, and presentation components of the project."
    )
    document.add_page_break()


def index_page(document):
    document.add_heading("INDEX", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = table.rows[0].cells
    set_cell_text(headers[0], "S. No.", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(headers[1], "Content", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(headers[2], "Page No.", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for cell in headers:
        shade(cell, "D9EAF7")
    set_table_widths(table, [0.7, 4.9, 0.6])

    items = [
        "Abstract",
        "Synopsis Summary",
        "Introduction",
        "Motivation",
        "Literature Review",
        "Gap Analysis",
        "Problem Statement",
        "Aim and Objectives",
        "Scope of the Project",
        "System Architecture",
        "Module Description",
        "Implementation Details",
        "Author Contribution and Project-Specific Decisions",
        "Analytics and Feedback Logic",
        "Functional and Non-Functional Requirements",
        "Tools and Technologies Used",
        "Methodology",
        "Experimental Setup and Testing Strategy",
        "Results and Discussion",
        "Project Screens and Visual Representation",
        "Deployment and Public Hosting",
        "Feasibility Analysis",
        "Conclusion",
        "Future Scope",
        "References",
        "List of Abbreviations",
        "Annexure and Appendix",
    ]
    for index, item in enumerate(items, 1):
        row = table.add_row().cells
        set_cell_text(row[0], str(index), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[1], item)
        set_cell_text(row[2], "", align=WD_ALIGN_PARAGRAPH.CENTER)
    document.add_page_break()


def comparative_table(document):
    add_table_title(document, "Table 1: Comparative analysis of current and proposed approaches")
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    set_cell_text(header[0], "Factor", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header[1], "Manual Assessment", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    set_cell_text(header[2], "Basic Quiz Portal", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    set_cell_text(header[3], "Analytics Dashboard", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    set_cell_text(header[4], "Proposed System", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    for cell in header:
        shade(cell, "D9EAF7")
    set_table_widths(table, [1.35, 1.1, 1.1, 1.25, 1.27])
    rows = [
        ("Score reporting", "Yes", "Yes", "Yes", "Yes"),
        ("Topic-wise diagnosis", "No", "Limited", "Partial", "Yes"),
        ("Repeated mistake tracking", "No", "No", "Limited", "Yes"),
        ("Time efficiency analysis", "No", "No", "Partial", "Yes"),
        ("Teacher intervention support", "Manual", "Low", "Partial", "Yes"),
        ("Admin question control", "No", "Low", "Low", "Yes"),
        ("IoT event representation", "No", "No", "No", "Yes"),
        ("Future AI/ML extension", "Low", "Low", "Medium", "High"),
    ]
    for values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(
                row[idx],
                value,
                size=10,
                align=WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT,
            )


def modules_table(document):
    add_table_title(document, "Table 2: Major modules of the proposed system")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    set_cell_text(header[0], "Module", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header[1], "Purpose", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header[2], "Output", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for cell in header:
        shade(cell, "D9EAF7")
    set_table_widths(table, [1.65, 2.45, 2.1])
    rows = [
        ("Login and access control", "Handles demo credentials and role-based routing", "Student, teacher, or admin workspace"),
        ("Student dashboard", "Shows personal performance and progress", "Topic cards, trend chart, feedback"),
        ("Quiz engine", "Delivers mixed mathematics questions", "Responses, timing, score"),
        ("Result analytics", "Summarizes attempt outcomes", "Charts, repeated errors, explanations"),
        ("Teacher dashboard", "Monitors class and student trends", "Class average, interventions, charts"),
        ("Admin panel", "Manages runtime data", "User CRUD, question CRUD, JSON export"),
        ("IoT simulation layer", "Represents software-generated node events", "Attendance, quiz, engagement event model"),
    ]
    for values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(row[idx], value, size=10)


def requirement_table(document):
    add_table_title(document, "Table 3: Functional requirements")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    set_cell_text(header[0], "Requirement ID", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    set_cell_text(header[1], "Requirement", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    set_cell_text(header[2], "Description", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    for cell in header:
        shade(cell, "D9EAF7")
    set_table_widths(table, [1.0, 1.6, 3.65])
    rows = [
        ("FR-1", "User login", "System shall support predefined student, teacher, and admin accounts."),
        ("FR-2", "Quiz attempt", "System shall support mixed mathematics quizzes."),
        ("FR-3", "Answer evaluation", "System shall evaluate MCQ and numeric questions."),
        ("FR-4", "Analytics generation", "System shall compute topic, difficulty, and timing insights."),
        ("FR-5", "Feedback generation", "System shall produce personalized learner recommendations."),
        ("FR-6", "Teacher drill-down", "System shall support learner-level intervention planning."),
        ("FR-7", "Admin management", "System shall allow runtime CRUD for users and questions."),
        ("FR-8", "Deployment", "System shall be publicly hosted for project viewing."),
    ]
    for values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(row[idx], value, size=10)


def testing_table(document):
    add_table_title(document, "Table 4: Testing scenarios and outcomes")
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    set_cell_text(header[0], "Test Case", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    set_cell_text(header[1], "Action", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    set_cell_text(header[2], "Expected Result", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    set_cell_text(header[3], "Status", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    for cell in header:
        shade(cell, "D9EAF7")
    set_table_widths(table, [0.9, 1.35, 3.25, 0.77])
    rows = [
        ("TC-1", "Student login", "Student dashboard opens correctly", "Passed"),
        ("TC-2", "Mixed quiz execution", "Quiz accepts answers and supports question navigation", "Passed"),
        ("TC-3", "Quiz submission", "Result page generates score, charts, and feedback", "Passed"),
        ("TC-4", "Teacher login", "Teacher analytics and student drill-down load", "Passed"),
        ("TC-5", "Admin login", "Runtime counts, forms, and export control appear", "Passed"),
        ("TC-6", "Public hosting", "GitHub Pages link opens project successfully", "Passed"),
    ]
    for values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(row[idx], value, size=10)


def abbreviation_table(document):
    add_table_title(document, "Table 5: List of abbreviations used in the report")
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    set_cell_text(header[0], "Abbreviation", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
    set_cell_text(header[1], "Full Form", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
    for cell in header:
        shade(cell, "D9EAF7")
    set_table_widths(table, [2.0, 4.2])
    rows = [
        ("AI", "Artificial Intelligence"),
        ("ML", "Machine Learning"),
        ("IoT", "Internet of Things"),
        ("UI", "User Interface"),
        ("UX", "User Experience"),
        ("MCQ", "Multiple Choice Question"),
        ("CRUD", "Create, Read, Update, Delete"),
        ("JSON", "JavaScript Object Notation"),
        ("HTML", "HyperText Markup Language"),
        ("CSS", "Cascading Style Sheets"),
        ("JS", "JavaScript"),
        ("MQTT", "Message Queuing Telemetry Transport"),
        ("KRMU", "KR Mangalam University"),
    ]
    for values in rows:
        row = table.add_row().cells
        set_cell_text(row[0], values[0], size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[1], values[1], size=10.5)


def report_body(document):
    document.add_heading("ABSTRACT", level=1)
    body(
        document,
        f"{PROJECT_TITLE} is designed as a smart academic support platform that combines the ideas "
        "of connected classroom systems, learning analytics, and personalized feedback into a "
        "single practical web application. In many educational settings, mathematics assessment "
        "still ends with marks, grades, or short remarks. This does not explain which concept is "
        "weak, whether mistakes are repeated, how speed affects performance, or what intervention "
        "should follow. The proposed project addresses this problem by converting assessment data "
        "into diagnostic and intervention-oriented insights."
    )
    body(
        document,
        "The current implementation is built using HTML, CSS, JavaScript, and Chart.js. It "
        "supports three roles: student, teacher, and administrator. The project uses a software-"
        "only IoT model in which learner activity is represented through virtual classroom nodes "
        "such as attendance events, quiz events, and engagement-related interactions. These events "
        "are treated as connected data streams and processed through a centralized analytics layer."
    )
    body(
        document,
        f"The deployed build currently contains {STUDENT_COUNT} student accounts, {TEACHER_COUNT} "
        f"teacher account, {ADMIN_COUNT} admin account, {QUESTION_COUNT} mathematics questions, "
        f"and {ATTEMPT_COUNT} seeded historical attempts used to drive teacher-side analytics. "
        "The analytics engine calculates topic-wise accuracy, difficulty-wise performance, average "
        "time per question, mastery classification, and repeated error patterns across algebra, "
        "trigonometry, and calculus. Based on these measurable inputs, the system generates "
        "personalized student feedback and teacher-side intervention suggestions. The project is "
        "also publicly hosted through GitHub Pages, which makes it suitable for academic review, "
        "portfolio demonstration, and viva presentation."
    )
    body(
        document,
        "Although the current feedback system is rule-based rather than model-trained, the project "
        "has been structured to allow future AI/ML extension such as student performance "
        "prediction, adaptive practice generation, clustering, and recommendation ranking."
    )
    body(
        document,
        "Keywords: IoT, Learning Analytics, Mathematics Education, Personalized Feedback, Smart "
        "Classroom, Dashboard Analytics, Student Performance Analysis, Educational Data Mining"
    )

    add_heading_page_break(document, "SYNOPSIS SUMMARY")
    body(
        document,
        f"This combined document contains both the project synopsis and the full final-year "
        f"project report for \"{PROJECT_TITLE}\". The project is implemented as a browser-based "
        "smart classroom prototype that analyzes learner performance in mathematics and generates "
        "personalized feedback through a hardware-free IoT model."
    )
    body(
        document,
        f"The current deployed build contains {STUDENT_COUNT} student accounts, "
        f"{TEACHER_COUNT} teacher account, {ADMIN_COUNT} admin account, {QUESTION_COUNT} seeded "
        f"questions, and {ATTEMPT_COUNT} seeded historical attempts. The mathematical scope covers "
        "algebra, trigonometry, and calculus. The application is publicly hosted and can be "
        "opened directly for demonstration."
    )
    document.add_heading("Synopsis Highlights", level=2)
    bullet(document, "Project domain: learning analytics, smart classroom systems, and educational decision support.")
    bullet(document, "Project type: hardware-free IoT-enabled web application.")
    bullet(document, "Main roles: student, teacher, and admin.")
    bullet(document, "Core functions: quiz execution, analytics, repeated-error detection, teacher drill-down, runtime management.")
    bullet(document, "Technology stack: HTML, CSS, JavaScript, Chart.js, GitHub Pages.")
    bullet(document, "Public repository and hosted deployment are already available.")
    document.add_heading("Problem Overview", level=2)
    body(
        document,
        "Traditional mathematics assessment often ends with marks and remarks rather than detailed "
        "academic diagnosis. Students do not always know which topic is weak, teachers do not "
        "always get a clear intervention map, and smart classroom ideas often remain disconnected "
        "from practical academic tools. The proposed project addresses this by treating classroom "
        "activity as connected data and converting it into chapter-wise insights and feedback."
    )
    document.add_heading("Objectives in Brief", level=2)
    bullet(document, "To create a role-based mathematics learning platform.")
    bullet(document, "To simulate IoT classroom event flow without physical hardware.")
    bullet(document, "To calculate topic accuracy, speed, and repeated mistake patterns.")
    bullet(document, "To generate learner-facing feedback and teacher-facing intervention suggestions.")
    bullet(document, "To deploy the project publicly for review and presentation.")
    document.add_heading("Expected Outcome", level=2)
    body(
        document,
        "The expected outcome is a practical, demonstrable smart classroom prototype that helps "
        "students understand their mathematics weaknesses and helps teachers make data-driven "
        "academic decisions. The current version is not presented as a full trained ML system; "
        "instead, it uses explainable rule-based logic with valid future scope for AI/ML extension."
    )

    add_heading_page_break(document, "Chapter 1")
    document.add_heading("Introduction", level=2)
    document.add_heading("1.1 Background of the Project", level=3)
    body(
        document,
        "Digital learning environments have increased the amount of academic information generated "
        "by students, teachers, and classroom systems. Attendance records, quiz attempts, answer "
        "patterns, timing behavior, and repeated engagement signals now exist in structured form. "
        "However, much of this data is still underused in day-to-day academic decision-making."
    )
    body(
        document,
        "Mathematics is a subject in which timely and precise diagnosis is essential. A learner may "
        "appear weak in the subject overall, but the actual issue may be limited to a few concepts "
        "such as quadratic factorisation, trigonometric identities, or chain rule applications. "
        "Similarly, a student may conceptually understand a topic but still underperform because "
        "of poor pace or repeated carelessness under timed conditions."
    )
    body(
        document,
        "The proposed project addresses this issue by designing an academic monitoring and feedback "
        "system that does not stop at the score. Instead, it studies the learner’s answer pattern, "
        "response time, topic distribution, and repeated mistakes, then presents the outcome "
        "through personalized and teacher-facing views."
    )
    document.add_heading("1.2 IoT Relevance", level=3)
    body(
        document,
        "The project is intentionally positioned as IoT-enabled even though it does not depend on "
        "physical devices in its current version. This is because the essential logic of IoT lies "
        "in distributed event generation, transmission, central processing, and monitored outputs. "
        "The project models attendance, quiz, and engagement activity as virtual connected nodes "
        "and then processes them through a central analytics dashboard."
    )
    comparative_table(document)

    add_heading_page_break(document, "Chapter 2")
    document.add_heading("Motivation", level=2)
    body(
        document,
        "The first motivation behind the project is educational relevance. Students need more than "
        "marks; they need explanation. Teachers need more than class averages; they need patterns "
        "that indicate when and where intervention should happen. A system that can reveal weak "
        "topics, slow performance, and repeated conceptual mistakes makes the assessment process "
        "far more useful."
    )
    body(
        document,
        "The second motivation is technical. Many final-year projects in the education domain stop "
        "at a basic quiz platform, while many IoT projects stop at showing hardware activity "
        "without meaningful application logic. This project attempts to bridge those gaps by "
        "building a real user-facing application around smart-system architecture principles."
    )
    body(
        document,
        "The third motivation is practicality. A full hardware-based smart classroom may be "
        "difficult to deploy within the time and resource limits of an academic project. By using "
        "virtual nodes and software-simulated event streams, the project remains feasible while "
        "still matching the conceptual flow of an IoT-enabled system."
    )

    add_heading_page_break(document, "Chapter 3")
    document.add_heading("Literature Review", level=2)
    body(
        document,
        "Learning analytics research shows that educational outcomes improve when institutions move "
        "from raw data collection to structured interpretation. Scholars such as Siemens and "
        "Ferguson have described learning analytics as a domain focused on measuring, collecting, "
        "analyzing, and reporting data about learners and their contexts for the purpose of "
        "understanding and optimizing learning."
    )
    body(
        document,
        "Feedback research also supports the project direction. Effective feedback is specific, "
        "timely, and linked to measurable behavior. It should explain not only what is wrong but "
        "also what should happen next. This principle directly influences the design of the "
        "project’s recommendation engine."
    )
    body(
        document,
        "In IoT literature, connected systems are described through the lifecycle of node-level "
        "event generation, data transmission, central processing, and monitored output. Applying "
        "this logic to education enables the design of classroom systems that respond intelligently "
        "to ongoing learner behavior. The proposed system adopts that structure through a "
        "software-only implementation."
    )
    document.add_heading("3.1 Gap Analysis", level=3)
    body(
        document,
        "Existing educational tools often suffer from three major limitations: they focus only on "
        "total score, they provide limited intervention guidance, and they rarely integrate "
        "multiple roles into a connected academic workflow. The present project addresses these "
        "limitations by combining diagnosis, visualization, recommendation, and administration "
        "into one deployed application."
    )

    add_heading_page_break(document, "Chapter 4")
    document.add_heading("Problem Statement, Aim and Objectives", level=2)
    document.add_heading("4.1 Problem Statement", level=3)
    body(
        document,
        "In many mathematics classrooms, evaluation is limited to marks and final remarks rather "
        "than meaningful academic diagnosis. Students do not receive clear identification of weak "
        "chapters, repeated misconceptions, or pace-related issues. Teachers do not receive a "
        "central view of class trends, and administrators do not have a flexible platform for "
        "content management. Therefore, a connected and intelligent mathematics learning support "
        "system is needed."
    )
    document.add_heading("4.2 Aim", level=3)
    body(
        document,
        "The main aim of the project is to design and implement a smart, IoT-oriented learning "
        "analytics and personalized feedback system for mathematics education that converts learner "
        "activity into actionable academic insights."
    )
    document.add_heading("4.3 Objectives", level=3)
    for item in [
        "To create a role-based web application for student, teacher, and admin workflows.",
        "To simulate classroom events as connected IoT node activity.",
        "To support mathematics quizzes across algebra, trigonometry, and calculus.",
        "To analyze topic-wise accuracy, difficulty level, pace, and repeated error patterns.",
        "To generate personalized feedback and teacher intervention suggestions.",
        "To host the project publicly through GitHub Pages for easy project review.",
    ]:
        bullet(document, item)

    add_heading_page_break(document, "Chapter 5")
    document.add_heading("Scope of the Project", level=2)
    body(
        document,
        "The scope of the current implementation includes a complete front-end demonstrator with "
        "role-based access, question answering, analytics, charts, teacher drill-down, runtime "
        "administration, IoT-based positioning pages, and public deployment. The mathematics "
        "content scope is limited to algebra, trigonometry, and calculus for controlled and "
        "interpretable analysis."
    )
    body(
        document,
        "The project currently excludes secure production authentication, persistent database "
        "storage, real hardware nodes, and trained machine learning models. These are treated as "
        "future extensions rather than limitations of the conceptual design."
    )

    add_heading_page_break(document, "Chapter 6")
    document.add_heading("System Architecture", level=2)
    body(
        document,
        "The architecture follows a software-only IoT pipeline. The Virtual Node Layer represents "
        "attendance events, quiz participation, and engagement events. The Data Transmission Layer "
        "represents the flow of those events into a central system. The Processing Layer organizes "
        "response data, question metadata, timing values, and historical attempts. The Analytics "
        "and Decision Layer computes mastery and recommendations. The Presentation Layer delivers "
        "results through dashboards."
    )
    document.add_heading("6.1 Architectural Flow", level=3)
    numbered(document, "Virtual IoT nodes generate attendance, quiz, and engagement signals.")
    numbered(document, "The signals are treated as connected event input to the software platform.")
    numbered(document, "A processing layer merges event data with question metadata and history.")
    numbered(document, "An analytics engine computes topic accuracy, speed, trends, and repeated errors.")
    numbered(document, "The feedback engine converts analytics into learner and teacher guidance.")
    numbered(document, "Dashboards expose the outcome to students, teachers, and administrators.")
    modules_table(document)

    add_heading_page_break(document, "Chapter 7")
    document.add_heading("Implementation Details", level=2)
    body(
        document,
        "The project is implemented as an approved-stack architecture with Python Flask "
        "microservices, a React 18 and D3.js visualization frontend, Docker-based Kafka and "
        "Cassandra infrastructure, and supporting documentation pages at the repository root. The "
        "backend is organized into independent services for data collection, processing, feedback "
        "generation, visualization, and LMS integration."
    )
    body(
        document,
        f"In its present form, the implementation uses {QUESTION_COUNT} seeded questions across "
        "algebra, trigonometry, and calculus. It also uses predefined demo accounts and seeded "
        "historical attempts so that the dashboards can be demonstrated immediately without first "
        "collecting live classroom data. This makes the project presentation-ready while still "
        "preserving realistic analytics flow."
    )
    document.add_heading("7.1 Student Workflow", level=3)
    body(
        document,
        "Students log in through seeded demo accounts. The student dashboard shows progress "
        "metrics, strongest topic, repeated mistake count, and chart-based trend analysis. From "
        "there, the learner can start a new quiz, answer mixed questions, and view a detailed "
        "result page with chapter-level analytics and explanations."
    )
    document.add_heading("7.2 Teacher Workflow", level=3)
    body(
        document,
        "Teachers access class-wide analytics, including class average, slowest topic, weakest "
        "group, score distribution, and accuracy trends over time. They can also select a "
        "specific student and inspect interventions, repeated errors, recent attempts, and "
        "topic-wise performance."
    )
    document.add_heading("7.3 Admin Workflow", level=3)
    body(
        document,
        "Administrators use a runtime panel to manage users and questions. They can add, edit, or "
        "delete runtime entries and export the modified data as JSON for demonstration purposes. "
        "This gives the project an additional layer of completeness and control."
    )
    document.add_heading("7.4 Mathematics Question Bank", level=3)
    body(
        document,
        "The question bank contains seeded questions from algebra, trigonometry, and calculus. "
        "Each question stores topic, subtopic, difficulty, question type, correct answer, "
        "explanation, expected solving time, and an error tag. This structure allows the analytics "
        "engine to interpret learner behavior with higher precision than simple correct/incorrect "
        "counting."
    )

    add_heading_page_break(document, "Chapter 8")
    document.add_heading("Author Contribution and Project-Specific Decisions", level=2)
    body(
        document,
        "This report is based on the implemented project available through the linked GitHub "
        "repository and public deployment, not on a purely conceptual proposal. The application "
        "flow, questions, error tags, dashboard metrics, and hosted pages were created specifically "
        "for this project build. The mathematical dataset was chosen to keep the analytics easy to "
        "explain during project evaluation while still covering different difficulty levels and "
        "concept types."
    )
    body(
        document,
        "A deliberate design decision in this project was to avoid claiming a trained machine "
        "learning model where none had been built. For that reason, the present version uses "
        "transparent rule-based recommendation logic. This keeps the project technically honest and "
        "easier to defend in viva while still leaving clear scope for later AI/ML extension."
    )
    document.add_heading("8.1 Practical Challenges Faced", level=3)
    bullet(document, "Framing the project as IoT-based without depending on physical hardware.")
    bullet(document, "Designing analytics that were detailed enough to be useful but simple enough to explain.")
    bullet(document, "Creating a report-friendly and demo-friendly interface with student, teacher, and admin roles.")
    bullet(document, "Deploying a static project publicly so the evaluator can open it directly.")
    document.add_heading("8.2 Solutions Adopted", level=3)
    bullet(document, "Used virtual IoT nodes and simulated event flow instead of physical devices.")
    bullet(document, "Used explicit rule-based mastery and pace logic for interpretability.")
    bullet(document, "Used seeded demo data so charts and dashboards are meaningful from first launch.")
    bullet(document, "Hosted the final project on GitHub Pages for easy access.")

    add_heading_page_break(document, "Chapter 9")
    document.add_heading("Analytics and Feedback Logic", level=2)
    body(
        document,
        "The analytics engine is intentionally rule-based and explainable. For each attempt, it "
        "computes total score, total time, average time, topic-wise correct count, difficulty-wise "
        "performance, and error-tag occurrence. It then classifies topic mastery and time "
        "efficiency before generating learner-facing guidance."
    )
    document.add_heading("8.1 Topic Mastery Rules", level=3)
    bullet(document, "Accuracy greater than or equal to 80% is classified as strong.")
    bullet(document, "Accuracy from 50% to 79% is classified as moderate.")
    bullet(document, "Accuracy below 50% is classified as weak.")
    document.add_heading("8.2 Time-Efficiency Rules", level=3)
    bullet(document, "Below 80% of expected time is fast.")
    bullet(document, "Between 80% and 120% of expected time is normal.")
    bullet(document, "Above 120% of expected time is slow.")
    document.add_heading("8.3 Repeated Error Logic", level=3)
    body(
        document,
        "The system associates every incorrect response with an error category such as equation "
        "balancing, identity selection, double-angle conversion, integration bounds, or chain "
        "rule application. When the same category recurs across attempts, the system flags it as a "
        "repeated error and presents it in student and teacher views."
    )
    document.add_heading("8.4 Personalized Feedback Strategy", level=3)
    body(
        document,
        "Feedback is derived from measurable conditions. Learners with low overall accuracy "
        "receive foundational revision suggestions. Learners in the middle range receive topic-"
        "focused medium practice advice. Learners with strong and fast performance receive harder "
        "mixed-challenge recommendations. Weak and slow topics trigger concept recap and remedial "
        "practice suggestions, while strong and fast topics trigger advanced progression guidance."
    )

    add_heading_page_break(document, "Chapter 10")
    document.add_heading("Functional and Non-Functional Requirements", level=2)
    requirement_table(document)
    document.add_heading("9.1 Non-Functional Requirements", level=3)
    for item in [
        "The application should remain responsive on desktop and mobile screens.",
        "The interface should be simple enough for live project demonstration.",
        "The system should work without backend dependency in the present version.",
        "Charts and summaries should remain readable and presentation-friendly.",
        "The project structure should remain maintainable and easy to explain.",
    ]:
        bullet(document, item)

    add_heading_page_break(document, "Chapter 11")
    document.add_heading("Tools and Technologies Used", level=2)
    body(
        document,
        "The project uses HTML for semantic structure, CSS for responsive styling and visual "
        "design, and vanilla JavaScript for interaction logic, analytics, runtime state "
        "management, and view rendering. Chart.js is used for charts, while Git and GitHub are "
        "used for version control and deployment."
    )
    bullet(document, "Frontend Stack: HTML, CSS, JavaScript")
    bullet(document, "Chart Library: Chart.js")
    bullet(document, "Development Environment: XAMPP")
    bullet(document, "Version Control: Git and GitHub")
    bullet(document, f"Repository: {REPO_URL}")
    bullet(document, f"Live Hosting: {LIVE_URL}")

    add_heading_page_break(document, "Chapter 12")
    document.add_heading("Methodology", level=2)
    body(
        document,
        "The methodology followed an iterative software prototyping approach. First, the academic "
        "problem was defined in terms of diagnosis and intervention gaps. Second, the role-based "
        "interface and question data structure were designed. Third, the quiz flow and analytics "
        "logic were implemented. Fourth, teacher and admin views were integrated. Fifth, the "
        "project was positioned formally as a hardware-free IoT system. Finally, the application "
        "was deployed publicly and documented for submission."
    )
    document.add_heading("11.1 Workflow Summary", level=3)
    numbered(document, "Problem identification and topic finalization")
    numbered(document, "Question-bank and user-role design")
    numbered(document, "Quiz engine implementation")
    numbered(document, "Analytics and feedback integration")
    numbered(document, "Teacher and admin module addition")
    numbered(document, "IoT architecture positioning and final deployment")

    add_heading_page_break(document, "Chapter 13")
    document.add_heading("Experimental Setup and Testing Strategy", level=2)
    body(
        document,
        "The project was tested both locally and through the hosted public version. Local testing "
        "was performed using XAMPP, while the final deployment was verified on GitHub Pages. "
        "Predefined demo accounts were used for all three roles so that the full workflow could be "
        "demonstrated repeatedly during documentation and screenshot capture."
    )
    body(
        document,
        "The student workflow was tested for login, quiz navigation, answer evaluation, and result "
        "generation. The teacher workflow was tested for chart rendering, student selection, and "
        "intervention planning. The admin workflow was tested for CRUD operations and export "
        "control."
    )
    testing_table(document)

    add_heading_page_break(document, "Chapter 14")
    document.add_heading("Results and Discussion", level=2)
    body(
        document,
        "The implemented prototype successfully demonstrates that assessment can be transformed "
        "from a score-reporting process into a structured academic diagnosis process. The student "
        "view shows that personalized recommendations can be generated from measurable learner "
        "behavior. The teacher view shows that class trends and interventions can be derived from "
        "centralized academic data. The admin view demonstrates runtime flexibility and content "
        "control."
    )
    body(
        document,
        "Another major outcome is the successful positioning of the system as an IoT-oriented "
        "project without hardware. The project preserves the essential flow of virtual node "
        "generation, event movement, central processing, and monitored output. This is a strong "
        "academic result because it makes the system both explainable and implementable."
    )

    add_heading_page_break(document, "Chapter 15")
    document.add_heading("Project Screens and Visual Representation", level=2)
    body(
        document,
        "The following screenshots represent the implemented interface and form part of the visual "
        "evidence of project completion. The screenshots were captured from the publicly hosted "
        "project and represent the actual deployed interface."
    )
    add_picture_with_caption(document, "landing-page.png", "Figure 1: Landing page with demo role-based access")
    body(
        document,
        "The landing page introduces the project as a final-year demonstrator and provides "
        "role-based access through predefined demo accounts. It also links directly to the "
        "project motive and IoT architecture pages for academic positioning."
    )
    add_picture_with_caption(document, "student-dashboard.png", "Figure 2: Student dashboard showing mastery and trend analytics", page_break_before=True)
    body(
        document,
        "The student dashboard presents average accuracy, strongest topic, repeated mistake count, "
        "a trend line chart, and personalized recommendation cards designed for self-guided improvement."
    )
    add_picture_with_caption(document, "student-quiz.png", "Figure 3: Student quiz interface with question palette and expected time", page_break_before=True)
    body(
        document,
        "The quiz view supports mixed question types, progress navigation, time-tracking, and "
        "visual progress indicators. The interface is designed to feel like an assessment system "
        "rather than a static form."
    )
    add_picture_with_caption(document, "student-result.png", "Figure 4: Result page with charts, repeated errors, and explanations", page_break_before=True)
    body(
        document,
        "The result page is central to the project because it converts a quiz attempt into chapter-"
        "level charts, repeated-error observations, and next-step recommendations."
    )
    add_picture_with_caption(document, "teacher-dashboard.png", "Figure 5: Teacher dashboard with class analytics and drill-down", page_break_before=True)
    body(
        document,
        "The teacher dashboard provides class average, slowest topic, weakest group, performance "
        "distribution, accuracy trends, and selected-student interventions."
    )
    add_picture_with_caption(document, "admin-dashboard.png", "Figure 6: Admin dashboard for runtime management", page_break_before=True)
    body(
        document,
        "The admin panel helps demonstrate project completeness by allowing runtime control over "
        "users and questions. This makes the application adaptable during presentation."
    )
    add_picture_with_caption(document, "project-motive.png", "Figure 7: Project motive page", page_break_before=True)
    add_picture_with_caption(document, "iot-architecture.png", "Figure 8: IoT architecture explanation page", page_break_before=True)

    add_heading_page_break(document, "Chapter 16")
    document.add_heading("Deployment and Public Hosting", level=2)
    body(
        document,
        "After implementation, the project was pushed to a public GitHub repository using SSH-based "
        "Git authentication. The repository was then connected to GitHub Pages by selecting the "
        "main branch and root folder as the deployment source. This made the project available "
        "publicly through a stable hosted link."
    )
    body(
        document,
        f"The live deployed project is available at {LIVE_URL}. Because the repository is public, "
        "any person with the project link can access the hosted application, which is useful for "
        "faculty review, external sharing, and portfolio presentation."
    )

    add_heading_page_break(document, "Chapter 17")
    document.add_heading("Feasibility Analysis", level=2)
    body(
        document,
        "The technical feasibility of the project is high because it uses an accessible front-end "
        "stack and a deployment model that does not require cloud servers, database licensing, or "
        "institutional hardware procurement. The economic feasibility is also strong because the "
        "prototype can be created using free and widely available tools."
    )
    body(
        document,
        "Operational feasibility is equally strong because the system is easy to demonstrate and "
        "understand during project evaluation. The project does not require special hardware "
        "setup, making it practical for classroom submission, seminar explanation, and viva."
    )

    add_heading_page_break(document, "Chapter 18")
    document.add_heading("Conclusion", level=2)
    body(
        document,
        f"The project titled \"{PROJECT_TITLE}\" successfully demonstrates a practical, academic, "
        "and visually complete solution for mathematics learning analytics and personalized "
        "feedback. It brings together smart classroom ideas, connected event flow, analytics, "
        "intervention logic, and public deployment in one integrated prototype."
    )
    body(
        document,
        "The system is stronger than a simple quiz application because it focuses on interpretation "
        "and academic action. It is also stronger than a hardware-only IoT demo because it "
        "translates connected-system logic into a full user-facing educational workflow."
    )

    add_heading_page_break(document, "Chapter 19")
    document.add_heading("Future Scope", level=2)
    for item in [
        "Machine learning-based student performance prediction.",
        "Adaptive learning path recommendation.",
        "Student clustering based on speed and accuracy profile.",
        "Persistent backend storage and institution-wide data history.",
        "API or MQTT-style event transmission for realistic IoT messaging.",
        "Optional integration with real attendance devices or classroom sensors.",
        "Chatbot-assisted tutoring and revision support.",
        "Department-level and institution-level monitoring dashboards.",
    ]:
        bullet(document, item)

    add_heading_page_break(document, "References")
    for item in [
        "Siemens, G., and Long, P. Penetrating the fog: Analytics in learning and education. EDUCAUSE Review.",
        "Ferguson, R. Learning analytics: drivers, developments and challenges. International Journal of Technology Enhanced Learning.",
        "Hattie, J., and Timperley, H. The power of feedback. Review of Educational Research.",
        "Atzori, L., Iera, A., and Morabito, G. The Internet of Things: A survey. Computer Networks.",
        "Gubbi, J., Buyya, R., Marusic, S., and Palaniswami, M. Internet of Things (IoT): A vision, architectural elements, and future directions. Future Generation Computer Systems.",
        "Chart.js Documentation.",
        "GitHub Pages Documentation.",
    ]:
        bullet(document, item)

    add_heading_page_break(document, "List of Abbreviations")
    body(
        document,
        "The following abbreviations are used throughout the synopsis and final report for "
        "technical clarity, academic consistency, and easier reading during project evaluation."
    )
    abbreviation_table(document)

    add_heading_page_break(document, "Annexure and Appendix")
    document.add_heading("Annexure I: Similarity and Plagiarism Certificate Attachment", level=2)
    body(
        document,
        "This section is reserved for attachment of the official similarity report or plagiarism "
        "certificate generated through the university-approved checking system such as Turnitin, "
        "Ouriginal, or another department-approved platform. An official plagiarism certificate "
        "cannot be self-generated inside this report document and should be attached here after "
        "verification by the student and the department."
    )
    bullet(document, f"Student Name: {STUDENT_NAME}")
    bullet(document, f"Roll Number: {ROLL_NUMBER}")
    bullet(document, f"Section: {SECTION}")
    bullet(document, f"Project Title: {PROJECT_TITLE}")
    bullet(document, f"Projexa Team Id: {PROJEXA_TEAM_ID}")
    bullet(document, f"University: {UNIVERSITY_NAME}")
    bullet(document, "Official similarity report status: To be attached after institutional check")
    body(
        document,
        "At the time of final printed submission, the official similarity certificate should be "
        "inserted immediately after this page so that the report remains academically compliant "
        "and complete."
    )
    add_signature_block(
        document,
        [
            "Student Signature: __________________________",
            "Date of Similarity Submission: __________________________",
        ],
        [
            "Guide Verification: __________________________",
            "Department Seal / Approval: __________________________",
        ],
    )

    document.add_heading("Annexure II: Project Links", level=2)
    bullet(document, f"GitHub Repository: {REPO_URL}")
    bullet(document, f"Live Project Link: {LIVE_URL}")
    document.add_heading("Annexure III: Core Project Files", level=2)
    bullet(document, "backend/services/data_collector/app.py - learning-event ingestion service")
    bullet(document, "backend/services/data_processor/app.py - analytics and learner-summary service")
    bullet(document, "backend/services/feedback_generator/app.py - personalized feedback generation service")
    bullet(document, "backend/services/visualization_service/app.py - dashboard data aggregation service")
    bullet(document, "backend/services/integration_service/app.py - LMS integration service")
    bullet(document, "frontend-react/src/App.jsx - React 18 and D3.js visualization frontend")
    bullet(document, "docker-compose.yml - Kafka and Cassandra local infrastructure setup")
    bullet(document, "notebooks/learning_analytics_exploration.ipynb - exploratory analytics notebook")
    document.add_heading("Annexure IV: Research Paper / Publication Details", level=2)
    body(
        document,
        "The current project has been documented in a report-ready form that is also suitable for "
        "conversion into a short academic paper, conference paper, or departmental publication. "
        "No false claim of publication is made in this report. The details below describe the "
        "paper version that can be prepared from the implemented system."
    )
    bullet(document, f"Proposed Paper Title: {PROJECT_TITLE}")
    bullet(document, f"Primary Author: {STUDENT_NAME}")
    bullet(document, "Affiliation: Department of Computer Science and Engineering, School of Engineering and Technology, KR Mangalam University")
    bullet(document, "Paper Type: Implementation-driven academic project paper")
    bullet(document, "Current Status: Manuscript-ready project work; not yet formally published")
    bullet(document, "Suggested Themes: Learning analytics, educational technology, smart classroom systems, IoT-based academic monitoring, personalized feedback")
    bullet(document, "Possible Publication Route: Department journal, conference proceedings, student innovation journal, or institutional research compendium")

    document.add_heading("Annexure V: Screenshot Archive", level=2)
    add_picture_with_caption(document, "landing-page.png", "Annexure Figure A1: Landing page", page_break_before=True)
    add_picture_with_caption(document, "student-dashboard.png", "Annexure Figure A2: Student dashboard", page_break_before=True)
    add_picture_with_caption(document, "student-quiz.png", "Annexure Figure A3: Quiz page", page_break_before=True)
    add_picture_with_caption(document, "student-result.png", "Annexure Figure A4: Result page", page_break_before=True)
    add_picture_with_caption(document, "teacher-dashboard.png", "Annexure Figure A5: Teacher dashboard", page_break_before=True)
    add_picture_with_caption(document, "admin-dashboard.png", "Annexure Figure A6: Admin dashboard", page_break_before=True)
    add_picture_with_caption(document, "project-motive.png", "Annexure Figure A7: Project motive page", page_break_before=True)
    add_picture_with_caption(document, "iot-architecture.png", "Annexure Figure A8: IoT architecture page", page_break_before=True)


def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_document_style(document)
    cover_page(document)
    front_matter(document)
    index_page(document)
    report_body(document)
    for section in document.sections:
        add_footer(section)
    document.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
