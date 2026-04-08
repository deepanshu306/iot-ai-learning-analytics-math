from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = Path("output/doc/iot-ai-learning-analytics-project-report.docx")

PROJECT_TITLE = (
    "IoT-Enabled AI-Driven Learning Analytics and Personalized Feedback "
    "System for Mathematics Education"
)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_doc_defaults(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"

    document.styles["Title"].font.size = Pt(20)
    document.styles["Heading 1"].font.size = Pt(16)
    document.styles["Heading 2"].font.size = Pt(14)
    document.styles["Heading 3"].font.size = Pt(12)


def add_center_line(document, text="", bold=False, size=12, space_after=6):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(space_after)
    return para


def add_body_paragraph(document, text, bold=False):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return para


def add_bullet(document, text):
    para = document.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_cover_page(document):
    add_center_line(document, "Major Project Report", bold=True, size=16, space_after=18)
    add_center_line(document, PROJECT_TITLE, bold=True, size=18, space_after=16)
    add_center_line(document, "Project Category: University Based Project", size=12, space_after=10)
    add_center_line(
        document,
        "Submitted in partial fulfilment of the requirement of the degree of",
        size=12,
        space_after=4,
    )
    add_center_line(document, "BACHELOR OF TECHNOLOGY (CSE / AI & ML)", bold=True, size=14, space_after=4)
    add_center_line(document, "Semester VIII", size=12, space_after=4)
    add_center_line(document, "to", size=12, space_after=4)
    add_center_line(document, "[University Name]", bold=True, size=14, space_after=18)
    add_center_line(document, "by", size=12, space_after=8)
    add_center_line(document, "[Student Name]", bold=True, size=13, space_after=4)
    add_center_line(document, "[Roll Number]    Section - [Section]", size=12, space_after=20)
    add_center_line(document, "Under the supervision of", size=12, space_after=10)
    add_center_line(document, "[Internal Mentor Name]", bold=True, size=12, space_after=2)
    add_center_line(document, "Assistant Professor / Project Guide", size=11, space_after=18)
    add_center_line(document, "Department of Computer Science and Engineering", bold=True, size=12, space_after=2)
    add_center_line(document, "School of Engineering and Technology", bold=True, size=12, space_after=2)
    add_center_line(document, "[University Name], [City], India", size=12, space_after=2)
    add_center_line(document, "April 2026", size=12, space_after=0)
    document.add_page_break()


def add_declaration(document):
    document.add_heading("DECLARATION", level=1)
    add_body_paragraph(
        document,
        "I hereby declare that this project report entitled "
        f'"{PROJECT_TITLE}" is my original work and has not been submitted elsewhere '
        "for the award of any degree, diploma, or certificate. The information "
        "presented in this report is true and correct to the best of my knowledge.",
    )
    document.add_paragraph("\n\n")
    add_body_paragraph(document, "Signature: __________________________")
    add_body_paragraph(document, "Name: [Student Name]")
    add_body_paragraph(document, "Roll Number: [Roll Number]")
    add_body_paragraph(document, "Date: __________________________")
    document.add_page_break()


def add_certificate(document):
    document.add_heading("CERTIFICATE", level=1)
    add_body_paragraph(
        document,
        "This is to certify that the project entitled "
        f'"{PROJECT_TITLE}" submitted by [Student Name], [Roll Number], to '
        "[University Name] is a record of bonafide project work carried out under "
        "my supervision and guidance and is worthy of consideration for the partial "
        "fulfilment of the degree of Bachelor of Technology in Computer Science and "
        "Engineering / Artificial Intelligence and Machine Learning.",
    )
    document.add_paragraph("\n\n")
    add_body_paragraph(document, "Project Guide: __________________________")
    add_body_paragraph(document, "Name: [Internal Mentor Name]")
    add_body_paragraph(document, "Designation: Assistant Professor")
    add_body_paragraph(document, "Department: Computer Science and Engineering")
    add_body_paragraph(document, "Date: __________________________")
    document.add_page_break()


def add_index(document):
    document.add_heading("INDEX", level=1)
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "S. No."
    headers[1].text = "Content"
    headers[2].text = "Page No."
    for cell in headers:
        set_cell_shading(cell, "D9EAF7")

    contents = [
        "Abstract",
        "Introduction",
        "Motivation",
        "Literature Review / Comparative Work Evaluation",
        "Gap Analysis",
        "Problem Statement",
        "Objectives",
        "Tools / Platforms Used",
        "Methodology",
        "Experimental Setup",
        "Evaluation Metrics",
        "Results and Discussion",
        "Conclusion and Future Work",
        "References",
        "Annexure I: Screenshots",
        "Annexure II: Source Code Summary",
        "Annexure III: Future Enhancement Notes",
    ]

    for idx, item in enumerate(contents, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = item
        row[2].text = ""

    document.add_page_break()


def add_existing_systems_table(document):
    document.add_paragraph("Table 1: Existing Systems Comparison", style="Heading 3")
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    header[0].text = "Factor"
    header[1].text = "Evaluation Criteria"
    header[2].text = "Conventional Test Portal"
    header[3].text = "Basic Quiz Dashboard"
    header[4].text = "Proposed System"
    for cell in header:
        set_cell_shading(cell, "D9EAF7")

    rows = [
        ("Assessment Depth", "Only marks and final score", "Yes", "Partial", "No"),
        ("Topic-wise Analysis", "Accuracy by algebra/trigonometry/calculus", "No", "Partial", "Yes"),
        ("Speed Analysis", "Time per question and pace pattern", "No", "No", "Yes"),
        ("Mistake Tracking", "Repeated error-tag identification", "No", "Limited", "Yes"),
        ("Teacher Insights", "Student drill-down and class summaries", "Limited", "Partial", "Yes"),
        ("Feedback", "Personalized recommendations", "Static", "Basic", "Yes"),
        ("IoT Relevance", "Connected event-driven learning data", "No", "No", "Yes"),
        ("Adaptability", "Scope for ML and predictive extension", "Low", "Medium", "High"),
    ]

    for item in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(item):
            cells[idx].text = value


def add_report_body(document):
    document.add_heading("ABSTRACT", level=1)
    add_body_paragraph(
        document,
        "The rapid digitization of education has increased the need for intelligent systems that "
        "can go beyond score reporting and support meaningful academic intervention. In "
        "mathematics education, students often receive marks without understanding the exact "
        "conceptual areas in which they are weak, the types of mistakes they repeat, or the "
        "learning behaviors that affect their performance. This project proposes an IoT-enabled, "
        "AI-driven learning analytics and personalized feedback system for mathematics education "
        "that addresses these limitations through a software-based smart classroom model.",
    )
    add_body_paragraph(
        document,
        "The current prototype is implemented as a web-based application using HTML, CSS, "
        "JavaScript, and Chart.js. It supports student, teacher, and administrator roles. The "
        "system collects learner activity from software-simulated IoT nodes such as attendance "
        "events, quiz events, and engagement logs. These connected data points are processed to "
        "compute topic-wise accuracy, time efficiency, repeated mistake patterns, and overall "
        "performance trends across algebra, trigonometry, and calculus.",
    )
    add_body_paragraph(
        document,
        "Based on these analytics, the system generates personalized feedback for learners and "
        "summary insights for teachers. The proposed architecture demonstrates how a hardware-free "
        "IoT model can still preserve the essential workflow of connected event generation, "
        "centralized processing, analytics, and dashboard-based intervention. The project is also "
        "designed to support future extension toward machine learning-based prediction, adaptive "
        "recommendation, and real-device integration.",
    )
    add_body_paragraph(
        document,
        "Keywords: Learning Analytics, Personalized Feedback, Mathematics Education, Internet of "
        "Things, Smart Classroom, Topic Mastery, Academic Intervention, Dashboard Analytics",
    )

    document.add_page_break()
    document.add_heading("Chapter 1", level=1)
    document.add_heading("Introduction", level=2)
    document.add_heading("1.1 Background of the Project", level=3)
    add_body_paragraph(
        document,
        "Mathematics learning requires continuous monitoring, conceptual diagnosis, and timely "
        "feedback. However, in many academic settings, assessment is still limited to marks, "
        "grades, and brief teacher remarks. This makes it difficult for students to understand "
        "their topic-level weaknesses and for teachers to provide scalable, evidence-based "
        "intervention across an entire class. In addition, smart classroom environments demand "
        "systems that can process multiple forms of learner activity rather than relying only on "
        "final test scores.",
    )
    add_body_paragraph(
        document,
        "Learning analytics provides a way to interpret educational data in order to improve "
        "teaching and learning outcomes. When combined with an IoT-style data collection model, "
        "it becomes possible to treat classroom actions such as attendance, assessment events, and "
        "engagement signals as connected streams of academic data. These signals can then be "
        "processed to identify performance trends, repeated errors, and intervention needs in a "
        "timely and structured way.",
    )
    add_body_paragraph(
        document,
        "The proposed system is developed as a web-based smart classroom prototype focused on "
        "mathematics education. It integrates simulated IoT nodes, analytics logic, role-based "
        "dashboards, and personalized feedback generation. This makes the project suitable for a "
        "software-only academic environment while preserving the architectural flow of a connected "
        "IoT-enabled educational monitoring system.",
    )
    add_existing_systems_table(document)

    document.add_page_break()
    document.add_heading("Chapter 2", level=1)
    document.add_heading("Motivation", level=2)
    add_body_paragraph(
        document,
        "The motivation behind this project comes from the limitations of conventional assessment "
        "methods in mathematics education. Students often know their final score but do not know "
        "which chapter is weak, whether they are improving over time, or whether they are making "
        "conceptual errors or speed-related errors. Teachers, on the other hand, face difficulty "
        "in manually tracking such patterns for multiple learners.",
    )
    add_body_paragraph(
        document,
        "Another important motivation is the rise of data-driven educational support systems and "
        "the relevance of IoT concepts in smart classrooms. Even without physical hardware, the "
        "core idea of connected data generation, centralized monitoring, and responsive feedback "
        "can be implemented in software. This project therefore combines educational analytics, "
        "decision support, and IoT architecture in a practical final-year prototype.",
    )

    document.add_page_break()
    document.add_heading("Chapter 3", level=1)
    document.add_heading("Literature Review / Comparative Work Evaluation", level=2)
    add_body_paragraph(
        document,
        "Existing research in learning analytics emphasizes the value of using learner data to "
        "support academic success, track progress, and guide intervention. Studies in educational "
        "data mining and analytics have shown that meaningful insights can be derived from quiz "
        "results, interaction logs, and performance histories. Feedback research further "
        "demonstrates that timely, targeted guidance improves learning outcomes more effectively "
        "than generic remarks.",
    )
    add_body_paragraph(
        document,
        "At the same time, research on the Internet of Things highlights the importance of "
        "distributed data generation, event transmission, and centralized monitoring in connected "
        "systems. When these ideas are applied to education, they support the design of smart "
        "classroom platforms that can observe learner activity through connected endpoints and "
        "transform raw events into actionable academic insights.",
    )
    add_body_paragraph(
        document,
        "Most classroom quiz portals still provide limited analytics and very basic feedback. "
        "The present project extends beyond simple digital testing by combining topic-wise "
        "analysis, response timing, repeated error detection, and personalized intervention logic "
        "within one connected educational workflow.",
    )

    document.add_page_break()
    document.add_heading("Chapter 4", level=1)
    document.add_heading("Gap Analysis", level=2)
    add_body_paragraph(
        document,
        "A clear gap exists between ordinary online quiz systems and intelligent academic support "
        "systems. Most available systems focus on content delivery and final score display rather "
        "than diagnostic interpretation. They do not sufficiently connect multiple learning signals "
        "such as attendance, timing behavior, topic accuracy, and repeated mistakes into one "
        "central decision-making layer.",
    )
    add_body_paragraph(
        document,
        "Another gap lies in IoT-based educational prototypes, which are often dependent on "
        "physical devices, sensors, or institutional infrastructure. This creates implementation "
        "barriers for student projects. The proposed system addresses that gap by presenting a "
        "hardware-free IoT model based on virtual nodes and event simulation while preserving the "
        "data-flow principles of connected systems.",
    )

    document.add_page_break()
    document.add_heading("Chapter 5", level=1)
    document.add_heading("Problem Statement", level=2)
    add_body_paragraph(
        document,
        "In many mathematics classrooms, evaluation is limited to marks, percentages, and final "
        "remarks, while important learning signals remain untracked. Such evaluation does not "
        "clearly explain why a student is underperforming, which concept is weak, whether mistakes "
        "are repeated, how classroom participation changes, or how performance evolves across "
        "attempts. Teachers also face difficulty in manually identifying patterns across many "
        "learners, especially when quick intervention is needed.",
    )
    add_body_paragraph(
        document,
        "Therefore, there is a need for a connected and intelligent academic support system that "
        "can collect simulated classroom data, analyze mathematics performance in detail, and "
        "generate personalized feedback for improved learning outcomes.",
    )

    document.add_page_break()
    document.add_heading("Chapter 6", level=1)
    document.add_heading("Objectives", level=2)
    objectives = [
        "To design a web-based smart classroom prototype for mathematics education.",
        "To simulate IoT-based classroom nodes for attendance, quiz activity, and engagement events.",
        "To analyze learner performance using score, accuracy, response time, and topic-wise mastery.",
        "To identify repeated misconceptions and weak mathematical areas.",
        "To generate personalized feedback for students and decision-support insights for teachers.",
        "To create a hardware-free IoT model that remains practical for final-year implementation.",
        "To provide future extensibility toward machine learning-based prediction and adaptive recommendations.",
    ]
    for objective in objectives:
        add_bullet(document, objective)

    document.add_page_break()
    document.add_heading("Chapter 7", level=1)
    document.add_heading("Tools / Platforms Used", level=2)
    tools = [
        "HTML for page structure and content layout.",
        "CSS for visual styling, responsive design, and theme control.",
        "JavaScript for analytics logic, quiz flow, role-based dashboard behavior, and personalized feedback generation.",
        "Chart.js for data visualization through graphs and comparative analytics.",
        "Git and GitHub for version control, public repository hosting, and project collaboration.",
        "GitHub Pages for public deployment of the static project.",
        "XAMPP for local execution and browser-based demonstration during development.",
    ]
    for tool in tools:
        add_bullet(document, tool)

    document.add_page_break()
    document.add_heading("Chapter 8", level=1)
    document.add_heading("Methodology", level=2)
    add_body_paragraph(
        document,
        "The methodology of the proposed system follows a connected data-processing pipeline. "
        "First, learner-related events are generated through software-simulated IoT nodes. These "
        "nodes model attendance, quiz attempts, answer submission timing, and engagement-related "
        "signals. The generated data is routed into a central processing layer where it is "
        "normalized and combined with question metadata.",
    )
    add_body_paragraph(
        document,
        "Second, the analytics engine computes topic-wise accuracy, difficulty-wise performance, "
        "average response speed, repeated error tags, and mastery classification. Rule-based "
        "feedback logic then maps the computed signals to intervention messages, recommendations, "
        "and practice guidance. The results are finally exposed through separate student, teacher, "
        "and administrator views in the browser interface.",
    )
    add_body_paragraph(
        document,
        "The overall methodology can be summarized as: Virtual IoT Nodes -> Data Transmission "
        "Layer -> Processing and Analytics Engine -> Feedback and Dashboard Layer.",
    )

    document.add_page_break()
    document.add_heading("Chapter 9", level=1)
    document.add_heading("Experimental Setup", level=2)
    add_body_paragraph(
        document,
        "The prototype is executed as a static front-end web application. The environment uses a "
        "local XAMPP server for development and GitHub Pages for hosted deployment. The system "
        "contains seeded student, teacher, and administrator accounts. The question bank covers "
        "algebra, trigonometry, and calculus with a mix of multiple-choice and numeric-answer "
        "questions. Session-based data storage is used to simulate runtime activity without a "
        "persistent backend database.",
    )
    add_body_paragraph(
        document,
        "Teacher analytics use seeded historical attempts together with current session data. "
        "Graphs and summary panels are rendered in the browser using Chart.js. The project also "
        "includes separate pages for project motive and IoT architecture to support final-year "
        "project explanation and presentation.",
    )

    document.add_page_break()
    document.add_heading("Chapter 10", level=1)
    document.add_heading("Evaluation Metrics", level=2)
    metrics = [
        "Overall quiz score and percentage accuracy.",
        "Topic-wise accuracy for algebra, trigonometry, and calculus.",
        "Difficulty-wise performance across easy, medium, and hard questions.",
        "Average response time and time-efficiency classification.",
        "Repeated error-tag frequency for misconception tracking.",
        "Topic mastery classification as strong, moderate, or weak.",
        "Teacher-side class average and student drill-down indicators.",
    ]
    for metric in metrics:
        add_bullet(document, metric)

    document.add_page_break()
    document.add_heading("Chapter 11", level=1)
    document.add_heading("Results and Discussion", level=2)
    add_body_paragraph(
        document,
        "The implemented prototype successfully demonstrates a complete browser-based workflow for "
        "student assessment, teacher monitoring, and runtime administration. Students can log in, "
        "attempt mathematics quizzes, and receive targeted feedback generated from their "
        "performance signals. Teachers can observe class trends, topic mastery patterns, and "
        "student-specific intervention needs using chart-based analytics. Administrators can manage "
        "the question bank and demo accounts in runtime memory.",
    )
    add_body_paragraph(
        document,
        "From a project perspective, the system achieves the key objective of transforming "
        "assessment into diagnosis and intervention. It also demonstrates that IoT-based academic "
        "monitoring can be modeled in software using virtual nodes and connected event flows. "
        "Although the current feedback layer is rule-based, the project establishes a strong "
        "foundation for future AI and machine learning expansion.",
    )

    document.add_page_break()
    document.add_heading("Chapter 12", level=1)
    document.add_heading("Conclusion and Future Work", level=2)
    add_body_paragraph(
        document,
        "The proposed project presents a practical and academically relevant approach to smart "
        "mathematics education. By combining IoT-style data flow, learning analytics, and "
        "personalized feedback in one web-based system, the project addresses the limitations of "
        "traditional score-only assessment. It helps learners understand their conceptual "
        "weaknesses and supports teachers in taking data-driven academic decisions.",
    )
    add_body_paragraph(
        document,
        "Future work can extend the system with backend persistence, predictive performance "
        "models, adaptive quiz generation, recommendation engines, chatbot-based tutoring, MQTT or "
        "API-driven event transmission, and optional physical hardware integration when required. "
        "These enhancements would move the system from a strong prototype toward a production-ready "
        "smart classroom platform.",
    )

    document.add_page_break()
    document.add_heading("References", level=1)
    references = [
        "Siemens, G. and Long, P. Penetrating the fog: Analytics in learning and education. EDUCAUSE Review.",
        "Ferguson, R. Learning analytics: drivers, developments and challenges. International Journal of Technology Enhanced Learning.",
        "Hattie, J. and Timperley, H. The power of feedback. Review of Educational Research.",
        "Atzori, L., Iera, A., and Morabito, G. The Internet of Things: A survey. Computer Networks.",
        "Gubbi, J., Buyya, R., Marusic, S., and Palaniswami, M. Internet of Things (IoT): A vision, architectural elements, and future directions. Future Generation Computer Systems.",
    ]
    for item in references:
        add_bullet(document, item)

    document.add_page_break()
    document.add_heading("Annexure I: Screenshots", level=1)
    add_body_paragraph(document, "Attach screenshots of the landing page, student dashboard, quiz screen, result page, teacher dashboard, administrator dashboard, project motive page, and IoT architecture page.")

    document.add_heading("Annexure II: Source Code Summary", level=1)
    add_body_paragraph(document, "Include the important source files such as index.html, styles.css, app.js, and any supporting documentation or deployment notes.")

    document.add_heading("Annexure III: Future Enhancement Notes", level=1)
    add_body_paragraph(document, "Document possible future additions such as machine learning prediction, adaptive recommendation, backend storage, API-based event streaming, and optional hardware integration.")


def add_footer(document):
    section = document.sections[-1]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(PROJECT_TITLE)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)


def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    set_doc_defaults(document)
    add_cover_page(document)
    add_declaration(document)
    add_certificate(document)
    add_index(document)
    add_report_body(document)
    add_footer(document)
    document.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
