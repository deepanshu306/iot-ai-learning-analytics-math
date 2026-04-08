from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = Path("output/doc/Deepanshu-Jain-Project-Synopsis.docx")

STUDENT_NAME = "Deepanshu Jain"
ROLL_NUMBER = "2201730107"
SECTION = "B"
UNIVERSITY_NAME = "KR MANGALAM UNIVERSITY"
CITY = "Gurgaon"
MENTOR_NAME = "Gaurav"
PROJECT_CATEGORY = "University Based Project"
PROGRAM = "BACHELOR OF TECHNOLOGY (CSE / AI & ML)"
SEMESTER = "Semester VIII"
SUBMISSION_MONTH = "April 2026"

PROJECT_TITLE = (
    "IoT-Enabled AI-Driven Learning Analytics and Personalized Feedback "
    "System for Mathematics Education"
)
PROJECT_CODE = "Math Analytics Studio"
REPO_URL = "https://github.com/deepanshu306/iot-ai-learning-analytics-math"
LIVE_URL = "https://deepanshu306.github.io/iot-ai-learning-analytics-math/"
STUDENT_COUNT = 3
TEACHER_COUNT = 1
ADMIN_COUNT = 1
QUESTION_COUNT = 12
ATTEMPT_COUNT = 6


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_margins(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for style in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        document.styles[style].font.name = "Times New Roman"

    document.styles["Title"].font.size = Pt(20)
    document.styles["Heading 1"].font.size = Pt(16)
    document.styles["Heading 2"].font.size = Pt(14)
    document.styles["Heading 3"].font.size = Pt(12)


def center(document, text, bold=False, size=12, after=6):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(after)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def body(document, text):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def bullet(document, text):
    para = document.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def numbered(document, text):
    para = document.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def footer(section):
    para = section.footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"{PROJECT_TITLE} | Page ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    add_page_number(para)


def cover_page(document):
    center(document, "Major Project Synopsis", bold=True, size=18, after=16)
    center(document, PROJECT_TITLE, bold=True, size=18, after=18)
    center(document, f"Project Category: {PROJECT_CATEGORY}", size=12, after=10)
    center(document, "Submitted in partial fulfilment of the requirement of the degree of", size=12, after=4)
    center(document, PROGRAM, bold=True, size=14, after=4)
    center(document, SEMESTER, size=12, after=4)
    center(document, "to", size=12, after=4)
    center(document, UNIVERSITY_NAME, bold=True, size=14, after=16)
    center(document, "by", size=12, after=8)
    center(document, STUDENT_NAME, bold=True, size=13, after=4)
    center(document, f"{ROLL_NUMBER}    Section - {SECTION}", size=12, after=18)
    center(document, "Under the guidance of", size=12, after=8)
    center(document, MENTOR_NAME, bold=True, size=12, after=2)
    center(document, "Project Guide", size=11, after=18)
    center(document, "Department of Computer Science and Engineering", bold=True, size=12, after=2)
    center(document, "School of Engineering and Technology", bold=True, size=12, after=2)
    center(document, f"{UNIVERSITY_NAME}, {CITY}, India", size=12, after=2)
    center(document, SUBMISSION_MONTH, size=12, after=0)
    document.add_page_break()


def declaration(document):
    document.add_heading("DECLARATION", level=1)
    body(
        document,
        f"I hereby declare that this synopsis entitled \"{PROJECT_TITLE}\" is the result of my "
        "own study, design, and implementation work carried out during the final year of the "
        "Bachelor of Technology program. The work has been prepared for academic submission and "
        "has not been copied from any other previously submitted project in full or in part. "
        "All information, explanations, and references included in this synopsis are true to the "
        "best of my knowledge."
    )
    body(document, f"Name: {STUDENT_NAME}")
    body(document, f"Roll Number: {ROLL_NUMBER}")
    body(document, f"Section: {SECTION}")
    body(document, "Signature: __________________________")
    body(document, "Date: __________________________")
    document.add_page_break()


def certificate(document):
    document.add_heading("CERTIFICATE", level=1)
    body(
        document,
        f"This is to certify that the synopsis entitled \"{PROJECT_TITLE}\" prepared by "
        f"{STUDENT_NAME} ({ROLL_NUMBER}) is a bonafide academic work carried out under guidance "
        "for the partial fulfilment of the degree of Bachelor of Technology in Computer Science "
        "and Engineering / Artificial Intelligence and Machine Learning. The work presented in "
        "this synopsis is suitable for academic consideration and project evaluation."
    )
    body(document, f"Project Guide: {MENTOR_NAME}")
    body(document, f"Student Name: {STUDENT_NAME}")
    body(document, f"University: {UNIVERSITY_NAME}")
    body(document, "Signature of Guide: __________________________")
    body(document, "Date: __________________________")
    document.add_page_break()


def index_page(document):
    document.add_heading("INDEX", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    header[0].text = "S. No."
    header[1].text = "Content"
    header[2].text = "Page No."
    for cell in header:
        shade(cell, "D9EAF7")

    entries = [
        "Abstract",
        "Introduction",
        "Motivation",
        "Literature Review",
        "Gap Analysis",
        "Problem Statement",
        "Aim and Objectives",
        "Scope of the Project",
        "Proposed System Overview",
        "System Architecture",
        "Module Description",
        "Analytics and Feedback Logic",
        "Functional and Non-Functional Requirements",
        "Tools and Technologies Used",
        "Methodology",
        "Experimental Setup and Testing Strategy",
        "Results and Discussion",
        "Feasibility and Project Significance",
        "Conclusion",
        "Future Scope",
        "References",
        "Annexure",
    ]
    for index, entry in enumerate(entries, 1):
        row = table.add_row().cells
        row[0].text = str(index)
        row[1].text = entry
        row[2].text = ""
    document.add_page_break()


def comparison_table(document):
    document.add_paragraph("Table 1: Comparative analysis of existing and proposed systems", style="Heading 3")
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    header[0].text = "Factor"
    header[1].text = "Conventional Assessment"
    header[2].text = "Simple Quiz Portal"
    header[3].text = "Analytics Dashboard"
    header[4].text = "Proposed System"
    for cell in header:
        shade(cell, "D9EAF7")

    values = [
        ("Score reporting", "Yes", "Yes", "Yes", "Yes"),
        ("Topic-wise diagnosis", "No", "Limited", "Partial", "Yes"),
        ("Repeated error detection", "No", "No", "Limited", "Yes"),
        ("Teacher drill-down", "No", "No", "Partial", "Yes"),
        ("Timing analysis", "No", "No", "Partial", "Yes"),
        ("Personalized feedback", "Manual", "Static", "Basic", "Yes"),
        ("IoT-style event flow", "No", "No", "No", "Yes"),
        ("Admin question management", "No", "Limited", "Limited", "Yes"),
        ("Scope for AI/ML extension", "Low", "Low", "Medium", "High"),
    ]
    for row_values in values:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value


def modules_table(document):
    document.add_paragraph("Table 2: Major modules of the proposed system", style="Heading 3")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    header[0].text = "Module"
    header[1].text = "Purpose"
    header[2].text = "Key Outputs"
    for cell in header:
        shade(cell, "D9EAF7")

    rows = [
        ("Login and Access Control", "Provides role-wise access for student, teacher, and administrator", "Secure role entry into the correct dashboard"),
        ("Student Dashboard", "Shows progress trend, mastery cards, and latest recommendations", "Score summary, strongest topic, repeated mistake count"),
        ("Quiz Engine", "Serves mixed mathematics questions and records answers and timing", "Question responses, score, accuracy, time data"),
        ("Analytics Engine", "Processes attempt data to build topic and difficulty performance", "Mastery levels, pace classification, trend analysis"),
        ("Feedback Engine", "Converts analytics into personalized academic guidance", "Remedial suggestions and challenge recommendations"),
        ("Teacher Dashboard", "Supports class monitoring and student drill-down", "Class average, weakest group, slowest topic, interventions"),
        ("Admin Panel", "Manages runtime user and question data", "CRUD operations and JSON export"),
        ("IoT Simulation Layer", "Models attendance, quiz, and engagement as connected events", "Virtual node activity and centralized event flow"),
    ]
    for row_values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value


def requirements_table(document):
    document.add_paragraph("Table 3: Functional requirements", style="Heading 3")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    header[0].text = "Requirement ID"
    header[1].text = "Requirement"
    header[2].text = "Description"
    for cell in header:
        shade(cell, "D9EAF7")

    rows = [
        ("FR-1", "User login", "The system shall allow predefined student, teacher, and admin users to log in."),
        ("FR-2", "Quiz execution", "The system shall allow students to attempt mixed mathematics quizzes."),
        ("FR-3", "Answer evaluation", "The system shall evaluate both MCQ and numeric responses."),
        ("FR-4", "Performance analytics", "The system shall compute topic-wise, difficulty-wise, and speed-based analytics."),
        ("FR-5", "Feedback generation", "The system shall generate personalized recommendations based on measured performance."),
        ("FR-6", "Teacher insights", "The system shall provide class-level summaries and learner-specific interventions."),
        ("FR-7", "Admin control", "The system shall allow runtime management of users and questions."),
        ("FR-8", "Export support", "The system shall allow JSON export of runtime-managed data."),
    ]
    for row_values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value


def testing_table(document):
    document.add_paragraph("Table 4: Representative testing scenarios", style="Heading 3")
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    header[0].text = "Test Case"
    header[1].text = "Input / Action"
    header[2].text = "Expected Result"
    header[3].text = "Status"
    for cell in header:
        shade(cell, "D9EAF7")

    rows = [
        ("TC-1", "Student login with demo credentials", "Student dashboard opens successfully", "Passed"),
        ("TC-2", "Quiz submission with MCQ and numeric answers", "Score and result analytics are generated", "Passed"),
        ("TC-3", "Teacher login", "Teacher analytics and student drill-down are visible", "Passed"),
        ("TC-4", "Admin add user and question", "Runtime data updates correctly", "Passed"),
        ("TC-5", "Export runtime JSON", "Export file download is initiated", "Passed"),
        ("TC-6", "GitHub Pages deployment", "Public project link opens successfully", "Passed"),
    ]
    for row_values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value


def content(document):
    document.add_heading("ABSTRACT", level=1)
    body(
        document,
        f"{PROJECT_TITLE} is proposed as a smart academic support system that combines the ideas "
        "of Internet of Things, learning analytics, educational data interpretation, and "
        "personalized feedback within a single software-based platform. In many classroom "
        "environments, mathematics performance is still judged only by marks, percentages, and "
        "final remarks. Such evaluation methods do not adequately explain which mathematical "
        "concepts are weak, what types of mistakes are repeated, how response speed affects "
        "performance, or what intervention should be given next. The present project addresses "
        "these limitations through a multi-role web application that transforms assessment data "
        "into diagnostic insights."
    )
    body(
        document,
        "The current implementation is built using HTML, CSS, JavaScript, and Chart.js. It "
        "supports student, teacher, and administrator roles in a single-page interactive "
        f"environment. The current build contains {STUDENT_COUNT} student accounts, "
        f"{TEACHER_COUNT} teacher account, {ADMIN_COUNT} admin account, {QUESTION_COUNT} seeded "
        f"questions, and {ATTEMPT_COUNT} seeded historical attempts. Student activities such as "
        "quiz attempts, answer timing, and session interactions are treated as data produced by "
        "software-simulated IoT classroom nodes. These event streams are analyzed through a "
        "central processing and analytics layer that computes topic-wise accuracy, difficulty-"
        "based performance, average response speed, repeated error patterns, and mastery levels "
        "for algebra, trigonometry, and calculus."
    )
    body(
        document,
        "The system further converts these analytics into personalized feedback messages, teacher "
        "intervention suggestions, and class-level visual summaries. A runtime administration "
        "module allows user and question management for demonstration purposes. The project is "
        "hosted publicly through GitHub Pages, which makes it accessible for project review and "
        "presentation. Although the current system uses rule-based personalization instead of a "
        "trained machine learning model, its structure is intentionally designed to support future "
        "AI/ML expansion such as prediction, clustering, and adaptive quiz recommendation."
    )
    body(
        document,
        "Keywords: IoT, Learning Analytics, Personalized Feedback, Smart Classroom, Mathematics "
        "Education, Educational Data Mining, Dashboard Analytics, Student Performance Analysis"
    )
    document.add_page_break()

    document.add_heading("Chapter 1", level=1)
    document.add_heading("Introduction", level=2)
    document.add_heading("1.1 Background", level=3)
    body(
        document,
        "Education systems are rapidly moving toward digitized learning environments in which "
        "teaching, assessment, and engagement are increasingly mediated through software. In such "
        "environments, large amounts of academic data are generated every day. These include "
        "attendance records, assessment results, click-based interaction data, time spent on "
        "tasks, and repeated academic behaviors. However, in many institutions, only a very small "
        "portion of this data is actually interpreted for academic improvement."
    )
    body(
        document,
        "Mathematics is one of the subjects in which timely diagnosis is especially important. A "
        "student may score poorly not because of a complete lack of understanding, but because of "
        "specific weakness in a few subtopics, conceptual confusion in repeated question types, or "
        "slow response behavior under timed conditions. When teachers do not have clear analytics, "
        "such patterns remain hidden and intervention is delayed."
    )
    body(
        document,
        "Learning analytics provides a mechanism to transform raw learner behavior into academic "
        "insight. At the same time, IoT introduces the concept of distributed event generation and "
        "centralized monitoring. In this project, these two domains are combined through a "
        "software-only architecture in which virtual classroom nodes generate connected learner "
        "events that are interpreted through an analytics dashboard."
    )
    document.add_heading("1.2 Need for the Study", level=3)
    body(
        document,
        "There is a strong academic need for systems that help teachers move beyond score "
        "reporting and toward diagnosis, interpretation, and intervention. There is also a need "
        "for student projects that can demonstrate IoT concepts without requiring expensive or "
        "logistically difficult hardware procurement. The proposed system satisfies both needs by "
        "using a hardware-free IoT model and a mathematically focused analytics workflow."
    )
    document.add_heading("1.3 Project Domain", level=3)
    body(
        document,
        "The project lies at the intersection of web engineering, smart classroom architecture, "
        "learning analytics, educational data interpretation, and AI-oriented recommendation "
        "logic. It is particularly suitable for a Computer Science and Engineering / AI & ML "
        "student because it involves data-driven reasoning, role-based interface design, "
        "interactive dashboard development, and clear future scope for predictive intelligence."
    )
    comparison_table(document)
    document.add_page_break()

    document.add_heading("Chapter 2", level=1)
    document.add_heading("Motivation", level=2)
    body(
        document,
        "The core motivation of this project arises from the gap between what assessment systems "
        "normally display and what students and teachers actually need. Students need to know not "
        "only whether they passed or failed, but why they are weak, which topic requires revision, "
        "whether their speed is appropriate, and which misconception is occurring repeatedly. "
        "Teachers need a way to observe patterns across many learners without manually checking "
        "every response in detail."
    )
    body(
        document,
        "A second motivation comes from the idea of smart classrooms. Modern academic spaces are "
        "increasingly discussed in terms of connected monitoring, responsive systems, and "
        "data-driven educational support. However, many student projects stop at either dashboard "
        "visualization or isolated hardware demonstration. This project instead tries to preserve "
        "the logic of connected systems through virtual IoT nodes and an academic intervention "
        "engine."
    )
    body(
        document,
        "A third motivation is project practicality. Building a complete hardware-based IoT "
        "system for classroom monitoring may not be realistic within the time, cost, and approval "
        "constraints of a final-year submission. A software-simulated IoT model makes the project "
        "implementable while still allowing the student to demonstrate architecture, analytics, "
        "workflow, and design depth."
    )
    document.add_page_break()

    document.add_heading("Chapter 3", level=1)
    document.add_heading("Literature Review", level=2)
    body(
        document,
        "Research in learning analytics emphasizes that educational decisions become more effective "
        "when raw learner data is translated into measurable patterns and actionable academic "
        "insights. Authors such as Siemens and Ferguson have highlighted the value of analytics in "
        "monitoring progression, identifying risk, and improving teaching quality. In mathematics "
        "education, this is especially relevant because conceptual mastery depends on gradual, "
        "trackable progress across related topics."
    )
    body(
        document,
        "Feedback studies, including widely cited work by Hattie and Timperley, show that specific "
        "and timely feedback improves performance more effectively than generic remarks. The "
        "quality of the feedback matters: it should be tied to measurable academic behavior, "
        "targeted to the learner’s present level, and useful for planning the next learning step. "
        "The present project adopts this principle by generating recommendations from topic "
        "accuracy, pace, and repeated error patterns."
    )
    body(
        document,
        "IoT literature describes connected systems in terms of event generation, transmission, "
        "processing, and monitoring. Atzori and Gubbi discuss how IoT systems are fundamentally "
        "about the movement and interpretation of data produced by distributed sources. This idea "
        "is directly applicable to education when classroom actions are modeled as connected events. "
        "The current project borrows this architecture, but implements the node layer through "
        "software simulation rather than physical hardware."
    )
    body(
        document,
        "Existing quiz portals and classroom dashboards are often limited in one of two ways: "
        "either they provide only basic assessment summaries, or they provide visual charts without "
        "clear recommendation logic. The proposed system addresses this gap by adding rule-based "
        "personalization and teacher-specific intervention support to the analytics workflow."
    )
    document.add_page_break()

    document.add_heading("Chapter 4", level=1)
    document.add_heading("Gap Analysis", level=2)
    body(
        document,
        "The first gap identified in current systems is the lack of diagnostic depth. Traditional "
        "systems usually report total marks, percentages, and pass/fail status. They rarely provide "
        "subtopic-level insight, error repetition analysis, or time-efficiency interpretation. As "
        "a result, the assessment outcome remains descriptive instead of diagnostic."
    )
    body(
        document,
        "The second gap is the absence of integrated role-specific decision support. Students, "
        "teachers, and administrators each require different kinds of insight. Students need "
        "personalized guidance; teachers need class-level trends and intervention ideas; "
        "administrators need content and user management. Many systems address one of these views "
        "but not all within one coherent design."
    )
    body(
        document,
        "The third gap is the difficulty of implementing IoT-based academic projects in a feasible "
        "student environment. Hardware-centric prototypes may be expensive, institution-dependent, "
        "or incomplete. A software-only IoT architecture fills this gap by preserving the event "
        "pipeline without overcommitting to hardware procurement."
    )
    document.add_page_break()

    document.add_heading("Chapter 5", level=1)
    document.add_heading("Problem Statement", level=2)
    body(
        document,
        "In many mathematics classrooms, assessment is restricted to score reporting rather than "
        "meaningful academic diagnosis. The absence of topic-level analytics, repeated error "
        "tracking, pace analysis, and intervention-oriented feedback limits both student "
        "self-improvement and teacher decision-making. Furthermore, smart classroom systems are "
        "often discussed in theory but not implemented in a practical, software-driven academic "
        "prototype. Therefore, there is a need to develop a connected, intelligent, and accessible "
        "system that can collect learner activity, analyze mathematics performance, and generate "
        "personalized academic support through an IoT-oriented architecture."
    )
    document.add_page_break()

    document.add_heading("Chapter 6", level=1)
    document.add_heading("Aim and Objectives", level=2)
    body(document, "The main aim of the project is to develop a smart mathematics learning support system that combines IoT-style connected data flow, learning analytics, and personalized feedback generation in a single practical web application.")
    document.add_heading("6.1 Objectives", level=3)
    for text in [
        "To build a web-based prototype for mathematics performance monitoring.",
        "To simulate IoT classroom nodes for attendance, quiz, and engagement events.",
        "To provide student, teacher, and administrator roles with separate dashboards.",
        "To evaluate quiz responses using both multiple-choice and numeric-answer logic.",
        "To compute topic-wise accuracy across algebra, trigonometry, and calculus.",
        "To classify performance by difficulty level and time efficiency.",
        "To detect repeated error patterns from recent and historical attempts.",
        "To generate personalized feedback based on measurable learner behavior.",
        "To support teachers with class summaries and learner-specific interventions.",
        "To host the project publicly for demonstration and academic presentation.",
    ]:
        bullet(document, text)
    document.add_page_break()

    document.add_heading("Chapter 7", level=1)
    document.add_heading("Scope of the Project", level=2)
    body(
        document,
        "The project scope includes a complete front-end prototype capable of demonstrating the "
        "major academic and system-level ideas of the proposed solution. It includes login access, "
        "quiz execution, analytics calculation, role-based dashboards, teacher monitoring, admin "
        "data management, and public deployment. The mathematical scope currently focuses on "
        "algebra, trigonometry, and calculus, which together provide sufficient variation for topic "
        "analytics and personalized feedback generation."
    )
    body(
        document,
        "The present scope intentionally excludes persistent database storage, secure production "
        "authentication, trained machine learning models, and physical sensor integration. These "
        "areas are treated as future extensions. The objective of the current version is to show a "
        "fully functional and academically strong prototype rather than a complete enterprise "
        "system."
    )
    document.add_page_break()

    document.add_heading("Chapter 8", level=1)
    document.add_heading("Proposed System Overview", level=2)
    body(
        document,
        "The proposed system is a role-driven smart classroom application in which educational "
        "data is treated as connected event flow. Student-side activity begins with login and quiz "
        "execution. Every question response produces data related to correctness, topic, "
        "difficulty, subtopic, error category, and time spent. The analytics layer then organizes "
        "this data into topic mastery statistics, time-efficiency categories, score trends, and "
        "repeated misconception patterns."
    )
    body(
        document,
        "The student dashboard focuses on self-improvement, the teacher dashboard focuses on "
        "monitoring and intervention, and the admin panel focuses on managing runtime system data. "
        "The project also includes separate explanatory pages for project motive and IoT "
        "architecture so that the academic positioning is clear during project presentation."
    )
    document.add_page_break()

    document.add_heading("Chapter 9", level=1)
    document.add_heading("System Architecture", level=2)
    body(
        document,
        "The architecture of the proposed system follows a connected pipeline model. The first "
        "layer is the Virtual Node Layer, which simulates attendance events, quiz events, and "
        "engagement signals. The second layer is the Data Transmission Layer, which treats these "
        "events as connected messages entering the central system. The third layer is the "
        "Processing and Analytics Layer, where the raw data is normalized and evaluated. The final "
        "layer is the Presentation and Decision Layer, where the analyzed information is delivered "
        "through dashboards and personalized feedback."
    )
    body(
        document,
        "This design is important because it allows the project to remain technically grounded in "
        "IoT concepts without depending on physical hardware. The architecture mirrors the core "
        "idea of smart systems: multiple distributed sources generate data, a central engine "
        "interprets the data, and the results are presented to decision-makers."
    )
    document.add_heading("9.1 Architectural Flow", level=3)
    numbered(document, "Virtual IoT Nodes generate attendance, quiz, and engagement events.")
    numbered(document, "A centralized software layer receives and organizes these events.")
    numbered(document, "The analytics engine evaluates accuracy, pace, topic weakness, and repeated errors.")
    numbered(document, "The feedback engine maps analytics to recommendations and interventions.")
    numbered(document, "Dashboards present the result to students, teachers, and administrators.")
    document.add_page_break()

    document.add_heading("Chapter 10", level=1)
    document.add_heading("Module Description", level=2)
    modules_table(document)
    body(
        document,
        "Each module contributes to a specific part of the academic workflow. The student module is "
        "designed for visibility and self-guided improvement. The teacher module concentrates on "
        "monitoring and intervention. The admin module ensures that the application remains "
        "demonstrable and modifiable during project presentation. Together, these modules show "
        "both the software engineering depth and the educational logic of the project."
    )
    document.add_page_break()

    document.add_heading("Chapter 11", level=1)
    document.add_heading("Analytics and Feedback Logic", level=2)
    body(
        document,
        "The analytics engine of the project is based on measurable, explainable rules. For each "
        "quiz attempt, the system calculates total score, percentage accuracy, average time per "
        "question, topic-wise attempted count, topic-wise correct count, difficulty-wise "
        "performance, and error-tag frequency. This ensures that every recommendation generated by "
        "the system can be justified from visible learner behavior."
    )
    document.add_heading("11.1 Topic Mastery Logic", level=3)
    bullet(document, "Accuracy greater than or equal to 80% is classified as strong.")
    bullet(document, "Accuracy between 50% and 79% is classified as moderate.")
    bullet(document, "Accuracy below 50% is classified as weak.")
    document.add_heading("11.2 Time-Efficiency Logic", level=3)
    bullet(document, "Actual average time below 80% of expected time is treated as fast.")
    bullet(document, "Actual average time between 80% and 120% of expected time is treated as normal.")
    bullet(document, "Actual average time above 120% of expected time is treated as slow.")
    document.add_heading("11.3 Repeated Error Logic", level=3)
    body(
        document,
        "Each question is associated with an error tag such as equation balancing, quadratic "
        "factorisation, double-angle conversion, integration bounds, or chain rule. When the same "
        "error tag appears repeatedly across recent work, the system flags it as a repeated "
        "mistake pattern."
    )
    document.add_heading("11.4 Feedback Rules", level=3)
    bullet(document, "Low-performing learners receive foundational revision advice.")
    bullet(document, "Mid-range learners receive targeted topic revision and medium practice guidance.")
    bullet(document, "High-performing learners receive mixed hard-question challenge guidance.")
    bullet(document, "Weak and slow topics produce concept recap and remedial practice suggestions.")
    bullet(document, "Strong and fast topics produce challenge-set recommendations.")
    body(
        document,
        "Because the current system uses transparent rule-based logic, it remains technically "
        "honest while still appearing intelligent during academic demonstration. This is a strong "
        "position for a final-year project because it provides clear scope for future AI/ML "
        "extension without making false claims about trained model deployment."
    )
    document.add_page_break()

    document.add_heading("Chapter 12", level=1)
    document.add_heading("Functional and Non-Functional Requirements", level=2)
    requirements_table(document)
    document.add_heading("12.1 Non-Functional Requirements", level=3)
    for text in [
        "The application should remain responsive on desktop and mobile screens.",
        "The dashboards should be visually clear and interpretable during project demonstration.",
        "The application should support browser-based deployment without backend dependency.",
        "The interface should be easy to use for all three roles.",
        "The project should remain maintainable through a flat file structure and simple JavaScript logic.",
    ]:
        bullet(document, text)
    document.add_page_break()

    document.add_heading("Chapter 13", level=1)
    document.add_heading("Tools and Technologies Used", level=2)
    body(
        document,
        "The project uses a lightweight front-end stack so that the implementation remains easy to "
        "demonstrate, edit, and host. The application is developed in HTML for structure, CSS for "
        "responsive visual design, and JavaScript for all business logic including authentication, "
        "quiz flow, analytics, chart rendering, and admin management. Chart.js is used to create "
        "line, bar, and doughnut charts for student and teacher dashboards."
    )
    body(
        document,
        "XAMPP is used for local hosting during development. Git and GitHub are used for version "
        "control, source-code backup, and collaborative presentation. GitHub Pages is used to host "
        "the final deployed version publicly. This makes the project easy to access and share "
        "during viva, internal review, and final evaluation."
    )
    bullet(document, "Frontend Languages: HTML, CSS, JavaScript")
    bullet(document, "Visualization Library: Chart.js")
    bullet(document, "Development Environment: XAMPP")
    bullet(document, "Version Control: Git and GitHub")
    bullet(document, f"Repository URL: {REPO_URL}")
    bullet(document, f"Live Project URL: {LIVE_URL}")
    document.add_page_break()

    document.add_heading("Chapter 14", level=1)
    document.add_heading("Methodology", level=2)
    body(
        document,
        "The project methodology follows an iterative prototyping approach. In the first phase, "
        "the academic problem was defined in terms of limited diagnosis, weak personalization, and "
        "the need for an IoT-oriented classroom model. In the second phase, the user roles, quiz "
        "logic, question structure, and analytics requirements were identified. In the third phase, "
        "the application interface and business logic were implemented as a browser-based prototype."
    )
    body(
        document,
        "The fourth phase involved integrating the analytics engine with the quiz and historical "
        "attempt structure. The fifth phase added feedback generation, teacher interventions, and "
        "admin runtime CRUD capability. The final phase focused on deployment, project motive "
        "positioning, IoT architecture explanation, and live hosting through GitHub Pages."
    )
    body(
        document,
        "One important project decision was to keep the recommendation engine rule-based in the "
        "current version instead of falsely presenting a trained machine learning model. This "
        "keeps the work technically honest while making the logic easy to explain during viva and "
        "leaving valid scope for future AI/ML enhancement."
    )
    document.add_heading("14.1 Workflow Steps", level=3)
    numbered(document, "Requirement gathering and topic finalization")
    numbered(document, "Role-based interface planning")
    numbered(document, "Mathematics question-bank design")
    numbered(document, "Quiz and evaluation logic implementation")
    numbered(document, "Analytics and recommendation integration")
    numbered(document, "IoT positioning through virtual node modeling")
    numbered(document, "Testing, redesign, and deployment")
    document.add_page_break()

    document.add_heading("Chapter 15", level=1)
    document.add_heading("Experimental Setup and Testing Strategy", level=2)
    body(
        document,
        "The project was tested as a web-based prototype using demo accounts for student, teacher, "
        "and administrator roles. Historical sample attempts were seeded into the application to "
        "make trend analysis and teacher drill-down meaningful from the first run. The application "
        "was then tested with live student quiz execution to validate score generation, result "
        "analysis, repeated error identification, and feedback rendering."
    )
    body(
        document,
        "Teacher-side testing verified class average calculation, topic comparison, score "
        "distribution, selected-student breakdown, and intervention suggestion generation. "
        "Administrator testing verified question and user runtime CRUD, validation messages, and "
        "JSON export behavior. Public deployment testing verified that the project was successfully "
        "hosted through GitHub Pages and accessible via a public link."
    )
    testing_table(document)
    document.add_page_break()

    document.add_heading("Chapter 16", level=1)
    document.add_heading("Results and Discussion", level=2)
    body(
        document,
        "The developed project successfully demonstrates that a static front-end application can "
        "still model the essential behavior of a smart academic monitoring system. Students are "
        "able to attempt quizzes, receive immediate evaluation, and review performance using "
        "topic-wise visual summaries and recommendation cards. Teachers are able to identify weak "
        "groups, compare topic-level performance, and inspect individual learners without needing "
        "manual cross-checking of all responses."
    )
    body(
        document,
        "One important outcome of the project is that it shifts the value of assessment from "
        "simple reporting to actionable intervention. Instead of stopping at the final score, the "
        "system asks what the score means, which topic caused the loss, whether timing was a "
        "factor, and what should happen next. This is a more educationally meaningful result than "
        "traditional reporting."
    )
    body(
        document,
        "Another important result is the successful academic framing of the project as IoT-based "
        "without physical hardware. The project demonstrates that the essence of IoT can be "
        "captured through distributed event generation, transmission logic, central analytics, and "
        "dashboard-driven monitoring. This makes the prototype feasible, explainable, and aligned "
        "with smart system design principles."
    )
    document.add_page_break()

    document.add_heading("Chapter 17", level=1)
    document.add_heading("Feasibility and Project Significance", level=2)
    body(
        document,
        "The technical feasibility of the project is high because it uses a simple, accessible "
        "technology stack and does not require costly infrastructure. The project can be run "
        "locally through XAMPP and deployed publicly through GitHub Pages. The operational "
        "feasibility is also strong because the system is easy to demonstrate during viva and can "
        "be explained clearly through role-based flows."
    )
    body(
        document,
        "The academic significance of the project lies in its combination of software engineering, "
        "educational analytics, IoT architecture concepts, and future AI/ML scope. It is not just "
        "a quiz app and not just a dashboard; it is an intervention-oriented learning support "
        "prototype. This makes it appropriate for a final-year project in CSE / AI & ML."
    )
    document.add_page_break()

    document.add_heading("Chapter 18", level=1)
    document.add_heading("Conclusion", level=2)
    body(
        document,
        f"The project titled \"{PROJECT_TITLE}\" successfully demonstrates a connected, data-driven "
        "approach to mathematics education. It combines role-based dashboards, quiz analytics, "
        "feedback generation, and IoT-style event modeling inside a practical browser application. "
        "The project is strong from both a technical and academic standpoint because it transforms "
        "assessment into diagnosis and intervention."
    )
    body(
        document,
        "The current implementation proves that even without hardware and backend complexity, a "
        "meaningful smart classroom prototype can be built and publicly deployed. It is suitable "
        "for project demonstration, report submission, and future enhancement."
    )
    document.add_page_break()

    document.add_heading("Chapter 19", level=1)
    document.add_heading("Future Scope", level=2)
    for text in [
        "Integration of backend storage for persistent student history.",
        "Machine learning-based student performance prediction.",
        "Adaptive quiz generation based on recent learner behavior.",
        "Student clustering based on performance and pace patterns.",
        "Recommendation ranking using ML models instead of static rules.",
        "API or MQTT-style transmission for a more realistic IoT communication layer.",
        "Optional physical device integration for attendance or classroom event capture.",
        "Institution-level dashboards and multi-classroom monitoring.",
    ]:
        bullet(document, text)
    document.add_page_break()

    document.add_heading("References", level=1)
    refs = [
        "Siemens, G., and Long, P. Penetrating the fog: Analytics in learning and education. EDUCAUSE Review.",
        "Ferguson, R. Learning analytics: drivers, developments and challenges. International Journal of Technology Enhanced Learning.",
        "Hattie, J., and Timperley, H. The power of feedback. Review of Educational Research.",
        "Atzori, L., Iera, A., and Morabito, G. The Internet of Things: A survey. Computer Networks.",
        "Gubbi, J., Buyya, R., Marusic, S., and Palaniswami, M. Internet of Things (IoT): A vision, architectural elements, and future directions. Future Generation Computer Systems.",
        "Open-source documentation for Chart.js, GitHub Pages, and browser-based JavaScript deployment concepts.",
    ]
    for item in refs:
        bullet(document, item)
    document.add_page_break()

    document.add_heading("Annexure", level=1)
    document.add_heading("Annexure I: Project Links", level=2)
    bullet(document, f"GitHub Repository: {REPO_URL}")
    bullet(document, f"Live Deployed Project: {LIVE_URL}")
    document.add_heading("Annexure II: Main Files", level=2)
    bullet(document, "backend/services/data_collector/app.py - data ingestion microservice")
    bullet(document, "backend/services/data_processor/app.py - analytics processing microservice")
    bullet(document, "backend/services/feedback_generator/app.py - personalized feedback microservice")
    bullet(document, "backend/services/integration_service/app.py - LMS integration microservice")
    bullet(document, "frontend-react/src/App.jsx - React 18 dashboard implementation")
    bullet(document, "docker-compose.yml - Kafka and Cassandra infrastructure configuration")
    document.add_heading("Annexure III: Suggested Screenshots for Submission", level=2)
    bullet(document, "Landing page with demo login")
    bullet(document, "Student dashboard with trend chart")
    bullet(document, "Quiz page during question attempt")
    bullet(document, "Result page with topic and difficulty charts")
    bullet(document, "Teacher dashboard with class analytics")
    bullet(document, "Admin panel with runtime management")
    bullet(document, "Project motive page")
    bullet(document, "IoT architecture page")


def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_margins(document)
    cover_page(document)
    declaration(document)
    certificate(document)
    index_page(document)
    content(document)
    for section in document.sections:
        footer(section)
    document.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
